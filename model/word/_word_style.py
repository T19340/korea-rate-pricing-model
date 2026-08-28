# -*- coding: utf-8 -*-
"""부서 제안서 양식 — Word 리포트 공통 서식.

_build_word.py 안에 모듈 수준으로 박혀 있던 서식 정의를 꺼낸 것이다. 회의 사후
검증판·전망판처럼 같은 양식의 문서를 여러 개 찍어야 해서 분리했다. 서식 값
자체는 _build_word.py(v16, 2026-08-10)와 한 글자도 다르지 않다.

    rp = Report(FIGS_DIR)
    rp.title_block("제목", "부제", "FI운용1부 · ...")
    rp.heading("1) 개요")
    rp.bullet("", "...")
    rp.arrow("결론 한 줄")
    rp.picture("fig02_fit_snapshot.png", "그림 2. ...")
    rp.data_table([(...), (...)], hl_rows=(1,))
    rp.save(OUT_PATH)
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── 폰트 해결 ────────────────────────────────────────────────────────
# 양식은 HY신명조·HY견고딕을 쓴다. 둘 다 상용 폰트라 설치돼 있지 않은 PC에서는
# Word가 임의 글꼴로 대체해 모양이 무너진다. 그래서 만들 때 설치 여부를 보고
# HY → Pretendard → 맑은 고딕 순으로 고른다.
# 레지스트리 값 이름은 시스템 폰트의 경우 영문이다(맑은 고딕 = "Malgun Gothic",
# 바탕 = "Batang & ..."). 별칭을 함께 검사한다.
_ALIAS = {"맑은 고딕": ("맑은 고딕", "Malgun Gothic"),
          "바탕": ("바탕", "Batang"),
          "HY신명조": ("HY신명조",),
          "HY견고딕": ("HY견고딕",),
          "Pretendard": ("Pretendard",)}


def _installed_fonts():
    names = set()
    try:
        import winreg
    except ImportError:
        return names
    key = r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"
    for hive in (winreg.HKEY_LOCAL_MACHINE, winreg.HKEY_CURRENT_USER):
        try:
            with winreg.OpenKey(hive, key) as k:
                for i in range(winreg.QueryInfoKey(k)[1]):
                    names.add(winreg.EnumValue(k, i)[0])
        except OSError:
            pass
    return names


_INSTALLED = _installed_fonts()


def pick_font(*candidates):
    """설치된 첫 후보를 돌려준다. 하나도 없으면 마지막 후보(최종 대체)."""
    for c in candidates:
        for alias in _ALIAS.get(c, (c,)):
            if any(alias in n for n in _INSTALLED):
                return c
    return candidates[-1]


BODY = pick_font("HY신명조", "Pretendard", "맑은 고딕")     # 본문
TITLE_FONT = pick_font("HY견고딕", "Pretendard", "맑은 고딕")  # 제목
CAPTION = pick_font("바탕", "Pretendard", "맑은 고딕")       # 캡션·부기

NAVY_TITLE = RGBColor(0x1F, 0x4E, 0x79)
BLUE_HEAD = RGBColor(0x61, 0x91, 0xCE)
BLUE_SOFT = RGBColor(0x93, 0xA8, 0xDD)
BLACK = RGBColor(0x10, 0x10, 0x10)
GRAY = RGBColor(0x6E, 0x6E, 0x6E)

CENTER = WD_ALIGN_PARAGRAPH.CENTER
RIGHT = WD_ALIGN_PARAGRAPH.RIGHT
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY


class Report:
    def __init__(self, figs_dir):
        self.FIGS = Path(figs_dir)
        self.doc = Document()

        sec = self.doc.sections[0]
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
        sec.top_margin, sec.bottom_margin = Cm(2.0), Cm(1.8)
        sec.left_margin, sec.right_margin = Cm(2.0), Cm(2.0)
        sec.header_distance, sec.footer_distance = Cm(1.8), Cm(1.35)

        normal = self.doc.styles["Normal"]
        normal.font.name = BODY
        normal.font.size = Pt(12)
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
        normal.paragraph_format.alignment = JUSTIFY
        normal.paragraph_format.space_after = Pt(2)
        normal.paragraph_format.line_spacing = 1.12

    # ---------- 기본 요소 ----------
    def set_font(self, run, name=None, size=12, bold=False, color=BLACK):
        name = name or BODY
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    def para(self, text="", size=12, bold=False, color=BLACK, font=None,
             align=JUSTIFY, space_after=2):
        p = self.doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            self.set_font(p.add_run(text), font, size, bold, color)
        else:
            self.set_font(p.add_run(" "), font, size, bold, color)
            p.runs[0].text = ""
        return p

    def spacer(self, size=11):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        self.set_font(p.add_run(""), BODY, size, True, BLUE_SOFT)
        return p

    def heading(self, text):
        self.spacer()
        p = self.para(text, size=13, bold=True, color=BLUE_HEAD)
        self.spacer()
        return p

    def arrow(self, text):
        return self.para("→ " + text, size=12)

    def bullet(self, label, text):
        return self.para(f"{label} {text}" if label else text)

    def page_break(self):
        self.doc.add_page_break()

    # ---------- 그림 ----------
    def picture(self, name, caption, width_cm=15.2):
        p = self.doc.add_paragraph()
        p.alignment = CENTER
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(self.FIGS / name), width=Cm(width_cm))
        self.para(caption, size=9, color=GRAY, font=CAPTION,
                  align=CENTER, space_after=8)

    # ---------- 표 ----------
    def cell_borders(self, tbl, edges=("bottom",), sz=12, color="1F4E79"):
        tblPr = tbl._tbl.tblPr
        borders = tblPr.makeelement(qn("w:tblBorders"), {})
        for edge in edges:
            el = tblPr.makeelement(qn(f"w:{edge}"), {})
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:color"), color)
            borders.append(el)
        tblPr.append(borders)

    def data_table(self, rows, font_size=10.5, hl_rows=()):
        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = t.cell(i, j)
                cp = cell.paragraphs[0]
                cp.alignment = CENTER
                self.set_font(cp.add_run(val), BODY, font_size,
                              bold=(i == 0), color=(NAVY_TITLE if i == 0 else BLACK))
                if i == 0 or i in hl_rows:
                    tcPr = cell._tc.get_or_add_tcPr()
                    sh = tcPr.makeelement(qn("w:shd"), {})
                    sh.set(qn("w:val"), "clear")
                    sh.set(qn("w:fill"), "E8EDF7" if i == 0 else "F4F6FB")
                    tcPr.append(sh)
        return t

    def tbl_note(self, text):
        self.para(text, size=9, color=GRAY, align=CENTER, space_after=6)

    # ---------- 제목 블록 ----------
    def title_block(self, title, subtitle, byline):
        tbl = self.doc.add_table(rows=2, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.cell_borders(tbl, edges=("bottom",), sz=12)
        self.set_font(tbl.cell(0, 0).paragraphs[0].add_run(title),
                      TITLE_FONT, 22, True, NAVY_TITLE)
        self.set_font(tbl.cell(1, 0).paragraphs[0].add_run(subtitle),
                      TITLE_FONT, 15, False, NAVY_TITLE)
        self.para(byline, size=9, color=GRAY, font=CAPTION,
                  align=RIGHT, space_after=6)

    def save(self, path):
        self.doc.save(str(path))
        return path
