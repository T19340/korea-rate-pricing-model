# -*- coding: utf-8 -*-
"""v8 리포트의 Word 판 — 부서 제안서 양식 + 논문식(흑백 세리프) 그림.

v2 개정: (1) 과거 사이클 백테스트를 정식 섹션으로, (2) 충격 분석(만기별 표 +
전달률 해석 + 표본외 검증) 추가, (3) 그림을 figs_paper 논문 스타일로 전면 교체.
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
FIGS = HERE.parent / "figs_paper"
OUT = HERE / "금통위_인상확률_계단함수모형.docx"

NAVY_TITLE = RGBColor(0x1F, 0x4E, 0x79)
BLUE_HEAD = RGBColor(0x61, 0x91, 0xCE)
BLUE_SOFT = RGBColor(0x93, 0xA8, 0xDD)
BLACK = RGBColor(0x10, 0x10, 0x10)
GRAY = RGBColor(0x6E, 0x6E, 0x6E)

doc = Document()

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin, sec.bottom_margin = Cm(2.0), Cm(1.8)
sec.left_margin, sec.right_margin = Cm(2.0), Cm(2.0)
sec.header_distance, sec.footer_distance = Cm(1.8), Cm(1.35)

normal = doc.styles["Normal"]
normal.font.name = "HY신명조"
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "HY신명조")
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(2)
normal.paragraph_format.line_spacing = 1.12


def set_font(run, name="HY신명조", size=12, bold=False, color=BLACK):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def para(text="", size=12, bold=False, color=BLACK, font="HY신명조",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_font(p.add_run(text), font, size, bold, color)
    else:
        set_font(p.add_run(" "), font, size, bold, color)
        p.runs[0].text = ""
    return p


def spacer(size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("")
    set_font(r, "HY신명조", size, True, BLUE_SOFT)
    return p


def heading(text):
    spacer()
    p = para(text, size=13, bold=True, color=BLUE_HEAD)
    spacer()
    return p


def arrow(text):
    return para("→ " + text, size=12)


def bullet(label, text):
    return para(f"{label} {text}" if label else text)


def picture(name, caption, width_cm=15.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(FIGS / name), width=Cm(width_cm))
    para(caption, size=9, color=GRAY, font="바탕",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)


def cell_borders(tbl, edges=("bottom",), sz=12, color="1F4E79"):
    tblPr = tbl._tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in edges:
        el = tblPr.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def data_table(rows, widths=None, font_size=10.5, hl_rows=()):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(cp.add_run(val), "HY신명조", font_size,
                     bold=(i == 0), color=(NAVY_TITLE if i == 0 else BLACK))
            if i == 0 or i in hl_rows:
                sh = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
                sh.set(qn("w:val"), "clear")
                sh.set(qn("w:fill"), "E8EDF7" if i == 0 else "F4F6FB")
                cell._tc.get_or_add_tcPr().append(sh)
    return t


def tbl_note(text):
    para(text, size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=6)


# ══════════════════════════════ 제목 블록
tbl = doc.add_table(rows=2, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_borders(tbl, edges=("bottom",), sz=12)
c1 = tbl.cell(0, 0).paragraphs[0]
set_font(c1.add_run("KOFR OIS를 통한 금통위 인상확률 측정"),
         "HY견고딕", 22, True, NAVY_TITLE)
c2 = tbl.cell(1, 0).paragraphs[0]
set_font(c2.add_run("2026년 8·10·11월 금통위 인상확률 추정, 만기별 충격 분석과 과거 사이클 검증"),
         "HY견고딕", 15, False, NAVY_TITLE)
para("FI운용1부 · 2026. 8. 10 · 데이터 기준 2026-08-07 (상세 검증·부록: report.html v16)",
     size=9, color=GRAY, font="바탕", align=WD_ALIGN_PARAGRAPH.RIGHT,
     space_after=6)

# ══════════════════════════════ 1) 개요
heading("1) 개요")
bullet("", "목표: 시장가격에서 2026년 8·10·11월 금통위의 25bp 인상확률을 직접 계산하고, "
           "인상·동결 각 시나리오에서 만기별 시장금리가 얼마나 움직일지 추정")
bullet("", "자료: KOFR OIS 1주~1년 호가(한국자금중개 MID, 2022.2~), 통안·국고 민평, 금투협 BMSI 서베이")
bullet("", "배경: 미국은 CME FedWatch가 회의별 확률을 실시간 제공하나 한국은 공개 도구 부재. "
           "시장내재 기대 추출 자체는 한국은행이 AFNS 모형으로 국회 보고에 쓰는 표준 절차")
para()
bullet("", "방법: 기준금리는 금통위 날에만 바뀌므로 미래 금리 경로를 계단으로 두고, "
           "만기별 OIS 호가에서 계단 높이(회의별 인상폭)를 역산")
bullet("", "검증: 2013~2026년 정례 금통위 124회 백테스트 + 충격 산술의 표본외 검증(2020~2026, 53회) "
           "+ 한국은행 방법론(AFNS) 자체 재현·대조")
para()
arrow("8/27 50% · 10/22 95% · 11/26 36% — 연말 내재 기준금리 3.20% (남은 세 회의에 약 1.8회 인상 반영)")

# ══════════════════════════════ 2) 방법론
heading("2) 방법론")
bullet("(1)", "계단함수 모형")
para()
para("OIS 고정금리는 계약 기간 하루짜리 금리(KOFR) 복리 평균과 교환 — 즉 호가는 미래 경로의 확률가중 평균")
para("경로 시작점은 1주 OIS 호가로 캘리브레이션(2.740%, OIS-KOFR 베이시스 흡수), "
     "1주 안에 회의가 있으면 최근 KOFR로 대체")
para("2개월 예시: 61일 내내 동결이면 고정금리 2.746%, 인상 확실하면 2.915%(회의 뒤 41일 비중만 반영) "
     "— 실제 호가 2.840%는 두 눈금의 약 56% 지점")
picture("fig01_concept.png",
        "그림 1. 계단함수의 원리 — 2개월 OIS 호가는 인상·동결 두 경로의 확률가중 평균")
para("만기 1개월~1년 호가 6개를 가중 최소자승(평활·수축 포함, 가우스-뉴턴 반복)으로 결합 적합, "
     "회의별 인상폭 Δ를 풀어 확률 = Δ/25bp로 환산")
para()
bullet("(2)", "AFNS·DNS-VAR 분해(한은 방법론 재현): 한국은행이 국회 보고에 쓰는 무재정 넬슨-시겔"
              "(AFNS, CDR 2011)을 칼만필터 최우추정으로 직접 재현(월말 무이표 164개, 적합 RMSE 4.2bp). "
              "백테스트·일별 재실행에는 같은 계열의 2단계 근사(DNS-VAR) 병행 — 실제 금리와 기대의 "
              "차이가 기간프리미엄이며, 계단함수 결과의 프리미엄 오염 점검용. 충격 분석(4장)의 "
              "구조 전파 곡선(적재함수)도 이 모형에서 공급")
bullet("(3)", "교차검증 렌즈: 통안 스트립(중위 프리미엄 차감), 국고 단기 스트립(누적치만), "
              "CD-IRS 스트립(실무 관행의 형식화·누적치만), BMSI 서베이")
bullet("(4)", "PCA 분해(부서 병행 재추정): 커브를 모형 공식이 아니라 데이터가 스스로 요약하게 하는 "
              "통계 기법 — 15개 만기의 일별 변화에서 수준·기울기·곡률 세 주성분을 뽑아(설명력 93%), "
              "금통위 서프라이즈(전일 컨센서스 대비)가 만기별로 얼마나 퍼지는지 재는 독립 잣대. "
              "역할 둘: ① 충격 분석(4장)의 실증 베타·DNS 전파를 제3의 방법으로 검증, "
              "② 당일에서 멈추는 본 모형과 달리 1주·1개월 뒤까지 충격 확산을 연장 관측")
para()
arrow("서로 다른 원리의 여섯 도구가 같은 방향이면 신뢰, 갈라지면 그 폭이 곧 불확실성의 크기")

# ══════════════════════════════ 3) 결과
heading("3) 결과 — 시장이 반영한 확률")
data_table([
    ("금통위", "내재 인상폭", "25bp 인상확률", "회의 후 내재 기준금리"),
    ("2026-08-27", "+12.5bp", "50%", "2.88%"),
    ("2026-10-22", "+23.8bp", "95%", "3.11%"),
    ("2026-11-26", "+9.0bp", "36%", "3.20%"),
    ("2027-01 (가정)", "+17.8bp", "71%", "3.38%"),
    ("2027-04 (가정)", "+16.1bp", "65%", "3.54%"),
    ("2027-07 (가정)", "+11.1bp", "45%", "3.65%"),
], hl_rows=(1,))
tbl_note("2027년 행은 일정 미공표에 따른 분기 노드 가정치로, 프리미엄이 섞여 액면 확률로 읽지 않음")
picture("fig02_fit_snapshot.png",
        "그림 2. 내재 정책금리 경로(a)와 곡선 적합(b) — 여섯 호가가 가중 RMSE 0.8bp로 설명됨")
bullet("", "교차검증: 연말 내재/기대 기준금리 — 통안 2.96%(하단 왜곡) < DNS-VAR 기대 3.01% < "
           "AFNS 기대 3.07%(한은 방법론 재현) < OIS 3.20% < 국고 3.27% < CD-IRS 3.34%(관행 잣대, "
           "조달 스프레드로 상단 왜곡 최대). 통안(강세 프리미엄)과 국고·IRS(상단 왜곡)가 OIS를 "
           "사이에 두는 배열 자체가 프리미엄 구조의 증거. CD-IRS는 백테스트에서 모든 국면 무정보 "
           "이하(Brier 0.38 대 0.35)라 연말 누적의 참고로만 사용")
para()
bullet("", "한국은행 방법론(AFNS) 재현: 한은이 국회에 보고한 분해(2022.11~2023.1월중 국고3년-기준금리 "
           "격차 축소 -136bp 중 순기대 -102bp)를 같은 창에서 -137bp/-115bp로 재현 — 방향·규모 일치. "
           "무재정 보정항은 3년 1bp·10년 8bp로 작아, 기대 경로는 AFNS와 DNS-VAR가 연말 기준 1bp 안에서 "
           "포개짐(3.07 대 3.01의 차이는 대부분 시작점 웨지 처리, 방법 자체의 추정 잡음 ±0.1~0.2%p)")
bullet("", "재현이 드러낸 한은 방법론의 한계 → 역할 분담: AFNS는 월말 커브의 매끄러운 요인 모형이라 "
           "금통위 일정 정보가 없고(회의별 배분 원리적 불가), 연말 기대의 추정 잡음이 ±0.1~0.2%p로 "
           "25bp 격자 한 칸보다 거칠며, 전체 표본 재추정 구조라 이벤트 다음 날을 읽지 못함 — "
           "이 세 자리를 OIS 계단함수가 메움(회의 날짜 내장·일별 호가로 익일 재계산). 거꾸로 1년 너머 "
           "기대 수준과 프리미엄 분리는 AFNS의 몫 — 경쟁이 아니라 시계(時界)별 역할 분담")
picture("fig12_afns_comparison.png",
        "그림 3. 한국은행 방법론(AFNS) 재현 — 기대 경로 앞 10개월(a)과 만기별 기간프리미엄(b). "
        "(a)의 10개월 너머는 모형의 평균회귀 성질이 지배해 절단, (b)의 3개월 음수는 시작점 웨지·적합 잔차(0 부근으로 읽음)")
arrow("프리미엄을 걷어낸 순수 기대는 한은 방법론 그대로 돌려도 연말 3.0~3.1% — "
      "OIS 3.20%는 프리미엄 포함 상단")
para()
bullet("", "누적과 배분의 성질이 다름 — 연말 누적치는 추세로 움직이고, 회의별 배분은 하루 단위로 진동")
picture("fig03_cumulative_allocation.png",
        "그림 4. 누적(a)은 추세, 배분(b)은 진동 — 배분은 수렴값과 방향으로만 읽는다")
arrow("누적치는 수준 그대로 믿고, 회의별 배분은 수렴값(50/95/36%)과 변화 방향으로만 읽을 것")

# ══════════════════════════════ 4) 충격 분석
heading("4) 충격 분석 — 인상·동결 시 만기별 금리 반응")
bullet("", "원리(Kuttner 2001): 가격을 움직이는 것은 결정이 아니라 서프라이즈 = 결정 - 사전 반영분. "
           "P=50%이므로 8/27은 인상 시 +12.4bp, 동결 시 -12.6bp로 드물게 대칭")
bullet("(1)", "DNS 계열(당일 반응) — 두 갈래: ① 실증 베타 — 2011년 이후 금통위 발표일의 만기별 "
              "금리 변화를 서프라이즈(통안 91일 변화)에 회귀(유효 67회) ② DNS 기울기충격 전파 — "
              "인상 서프라이즈를 곡선의 기울기 충격으로 번역해 만기별 적재함수로 전파")
para()
picture("fig07_passthrough.png",
        "그림 5. 만기별 전달률 — 서프라이즈 1bp당 이동(bp). 실제 이동폭은 전달률 × 서프라이즈")
para("그림 5의 세로축은 서프라이즈 1bp당 전달률(β). 8월 깜짝 인상이면 서프라이즈 +12.4bp이므로 "
     "만기별 이동폭 = β × 12.4bp — 아래 표가 그 환산 결과")
para()
data_table([
    ("만기", "인상 시 (실증)", "인상 시 (DNS)", "동결 시 (실증)", "동결 시 (DNS)"),
    ("통안 3개월", "+12.4bp", "+12.4bp", "-12.6bp", "-12.6bp"),
    ("통안 6개월", "+8.8bp", "+11.7bp", "-8.8bp", "-11.7bp"),
    ("통안 1년", "+6.3bp", "+10.2bp", "-6.4bp", "-10.3bp"),
    ("국고 2년*", "+3.2bp", "+8.1bp", "-3.2bp", "-8.1bp"),
    ("국고 3년", "+5.0bp", "+6.5bp", "-5.0bp", "-6.6bp"),
    ("국고 5년", "+5.8bp", "+4.5bp", "-5.9bp", "-4.6bp"),
    ("국고 10년", "+5.3bp", "+2.4bp", "-5.3bp", "-2.4bp"),
    ("국고 20년", "+5.4bp", "+1.2bp", "-5.4bp", "-1.2bp"),
    ("국고 30년", "+3.4bp", "+0.8bp", "-3.5bp", "-0.8bp"),
], font_size=10, hl_rows=(5, 7))
tbl_note("서프라이즈 ±12.4/-12.6bp 기준. * 국고 2년은 표본이 짧아(2013~) 신뢰도 낮음. "
         "실증 베타는 일간 변화 기준이라 회견·점도표 효과가 일부 혼입")
picture("fig08_scenario_curves.png",
        "그림 6. 8/27 두 갈래 시나리오 — 위(인상)·아래(동결) 곡선 사이가 당일 노출 폭")
para()
bullet("(2)", "PCA 계열(독립 검증 + 시간 전개) — 부서 병행 재추정: 국고 제로커브 15개 만기의 상관 "
              "PCA 요인(설명력 93%)을 전일 컨센서스(폴 평균) 자 서프라이즈에 회귀(|서프라이즈| "
              "5bp 이상 24회, 2012~2025). 서프라이즈 자가 달라 절대 크기 비교는 불가 → 프로파일"
              "(3개월=1)로 대조")
para()
para("당일 프로파일(그림 7a): 데이터 잣대 둘(실증·PCA)은 1년 이내 사실상 일치(6개월 0.70 vs "
     "0.65, 1년 0.51 vs 0.56)하고 DNS만 6개월~2년 배 구간에서 홀로 높음(적재함수 형태 탓). "
     "장기는 PCA 0.27(10년)이 실증 0.43과 DNS 0.19 사이 → 단기 산술은 데이터 잣대 기준 방법 "
     "불문, 장기 서술도 제3의 방법이 지지")
para("충격 그 후 한 달(그림 7b): 10년 반응이 당일 +2.1bp → 1주 +5.0bp → 1개월 +9.9bp(20년 "
     "+10.2bp), 3개월은 +7.9 → +9.3bp — 충격의 무게중심이 장기로 이동. 1개월 기준 10년이 "
     "3년(+6.9bp)보다 커 기계적 3/10년 플래트닝은 스티프닝 위험")
para("한계: 1주·1개월은 미국 금리·수급 등 동시 충격 미통제, 부트스트랩 95% 구간 대부분 0 포함"
     "(그림 7b 음영) → 인과 추정이 아니라 위험 시나리오의 중심값으로만 사용")
picture("fig13_pca_shock.png",
        "그림 7. PCA 병행 재추정 — 당일 프로파일 3잣대 대조(a)와 충격의 시간 전개(b, 음영은 1개월 95% 부트스트랩 구간)")
para()
para("국면별 대조(그림 8): 실증 프로파일은 국면마다 모양 자체가 바뀜 — 완만기(13~19.6) 전 만기 "
     "평행(≈1), 급속 인상기(21.8~23.1) 1년부터 음수(장기 역행), 완화·재인상기(24.10~, 현재 국면) "
     "3년 2.7배 증폭 혹. DNS는 국면 불변 곡선 하나, PCA 잣대는 값으로는 완만기만 유효"
     "(급속인상 4회는 부호만·완화기 3회는 계산 불가). 전표본 "
     "실증 0.43(10년)은 서프라이즈² 가중 혼합(가중치 절반 이상이 2022년) → 충격 표의 장기물은 "
     "국면 조건부·범위로 읽을 것")
picture("fig14_regime_profiles.png",
        "그림 8. 국면별 당일 프로파일(3개월=1) — 실증(●)은 국면마다 모양이 바뀌고, DNS(□)는 곡선 하나, PCA(△)는 완만기만 측정 가능")
arrow("단기물은 방향 베팅이 아니라 ±12bp의 변동성 노출 — 두 추정이 갈라지는 장기물은 "
      "결정 숫자보다 K-점도표·회견이 지배하고, 한 달 시계에서는 장기물로 확산")

# ══════════════════════════════ 5) 검증
heading("5) 검증 — 과거 사이클에서의 성적")
bullet("(1)", "실전 검증(2026년 7월): 서베이 발표일(7/14) 모형 74% vs 서베이 66%, "
              "회의 전일 모형 91% — 7/16 실제 인상")
picture("fig04_next_meeting.png",
        "그림 9. 다음 금통위 인상확률 — 모형(실선), 서베이(마름모), 실제 결정(상단 표식)")
bullet("(2)", "확률의 백테스트: 2013.2~2026.7 정례 금통위 124회, 회의 전일 데이터만으로 "
              "3분류 브라이어 채점(낮을수록 정확)")
para()
data_table([
    ("기간", "회의", "통안 스트립", "KOFR OIS", "BMSI 서베이", "무정보(빈도)"),
    ("2013~2016 인하기", "47", "0.22", "—", "0.15", "0.22"),
    ("2017~2019.6 인상기", "20", "0.06", "—", "0.03", "0.18"),
    ("2019.7~2021.7 인하기", "17", "0.31", "—", "0.16", "0.29"),
    ("2021.8~2023.1 인상기", "12", "0.77", "—*", "0.20", "0.28"),
    ("2023.2~2024.9 동결기", "13", "0.37", "—*", "0.03", "0.00"),
    ("2024.10~2026.5 인하기", "14", "0.33", "0.26", "0.19", "0.41"),
    ("2026.7 인상 재개", "1", "0.00", "0.02", "0.23", "0.00"),
    ("전체 (124회)", "124", "0.29", "—", "0.13", "0.35"),
], font_size=10, hl_rows=(4, 8))
tbl_note("* OIS는 2022~23년 호가가 지표성(고정 스프레드)이라 2024.10 이후만 유효 — 유효표본 15회에서 "
         "OIS 0.24, 통안 0.31, 서베이 0.20, 무정보 0.48. 무정보 = 채점 기간의 실현 빈도(사후 기준)")
para("전 기간 1등은 서베이(0.13) — 동결이 잦은 한국에서 동결에 극단적 확신을 부여하고 대체로 맞음. "
     "시장내재의 실패는 국면이 분명함: 2021~22 급속 인상기(0.77, 50bp 스텝이 25bp 격자를 깸)와 "
     "2023 동결기(0.37, 오지 않은 인하를 계속 반영)")
para("시장의 부가가치는 전환점: 2024-11 '깜짝 인하'를 통안은 전일 69% 반영(서베이 17%). "
     "다만 2025-01 깜짝 동결은 OIS 혼자 크게 틀림 — 한 도구만 보지 말 것")
picture("fig10_backtest_timeline.png",
        "그림 10. 13년의 채점표 — 회의 전일 내재 순확률과 실제 결정(상단 인상·하단 인하)")
arrow("총점은 서베이 우위 — 시장가격의 값어치는 둘이 갈라지는 순간의 방향 정보"
      "(인하 전환점 반복 우위, 인상 쪽은 엇갈림). 상시 병행하다 괴리를 신호로 쓸 것")
para()
bullet("(3)", "충격 산술의 표본외 검증: 4절의 베타를 2011~2019년으로만 추정해 "
              "2020.1~2026.7 정례 금통위 53회에 적용, 예측 변동 대 실제 당일 변동을 국면별 채점")
para()
data_table([
    ("국면 (회의 수)", "통안3M 상관", "적중", "국고3Y 상관", "적중", "국고10Y 상관", "적중"),
    ("완화·동결 2020.1~21.7 (13회)", "+0.71", "78%", "+0.52", "78%", "+0.32", "67%"),
    ("급속 인상 2021.8~23.1 (12회)", "+0.82", "67%", "-0.48", "25%", "-0.41", "25%"),
    ("동결 2023.2~24.9 (13회)", "+0.17", "50%", "+0.15", "50%", "+0.06", "50%"),
    ("인하 2024.10~26.5 (14회)", "+0.35", "62%", "+0.41", "69%", "+0.29", "62%"),
    ("전체 풀링 (53회)", "+0.72", "62%", "-0.21", "56%", "-0.20", "51%"),
], font_size=9, hl_rows=(2,))
tbl_note("부호 적중은 |서프라이즈| ≥ 3bp 회의 기준. 국면당 12~14회 소표본 — 방향으로 읽을 것")
picture("fig11_shock_validation.png",
        "그림 11. 표본외 성적 — 대각선(완벽한 예측)을 거스르는 점은 3년·10년 모두 급속 인상기(▲)에 집중")
para("단기물의 방향 신호는 전 국면 유효(상관 +0.35~+0.82). 장기물의 실패는 3년·10년 공히 급속 "
     "인상기 한 국면에 집중(10년: 완화 +0.32·인하 +0.29 vs 급속 -0.41) — 전체 풀링의 -0.21(3년)·"
     "-0.20(10년)은 국면 혼합이 만든 착시. 4장 국면 대조(그림 8)와 같은 무늬. 크기의 과대 추정은 "
     "전 국면 공통")
arrow("단기물 산술은 국면 불문 그대로 쓰고, 장기물은 \"지금이 어느 과거와 닮았는가\"를 먼저 정할 것 "
      "— 2011~19형 정상 사이클이면 8월 인상 시 3년물 +14bp, 2021~22형 소통 지배면 0 부근, "
      "4절 표의 +5bp는 그 사이의 평균")

# ══════════════════════════════ 6) 한계 및 활용
heading("6) 한계 및 활용")
bullet("", "한계: 25bp 격자 가정, 위험중립 확률(예보 아님), 브로커 호가(T+1 등 관행 근사 반영 시 "
           "8월 40%대 중반 — 실무 범위 40~60%)")
bullet("", "확률 대역: 호가 ±1틱 몬테카를로 90% 대역 기준 8월 ±2%p, 10월 ±6%p, 11월 ±11%p "
           "— 뒤 회의일수록 성기게 읽을 것")
para()
bullet("", "활용: ① 일별 확률·연말 내재 수준 모니터링(재실행 자동화) ② 금통위·CPI 등 이벤트 직후 "
           "\"어느 회의의 확률로 소화됐는지\" 분해 ③ 자기 뷰와 시장 확률의 간극이 클 때가 포지션 근거")
para()
arrow("결론: 10월 인상은 기정사실(95%), 진짜 논쟁은 8월(50%) — 8/27의 관전 변수는 결정 자체보다 "
      "K-점도표이며, 확률의 수준보다 변화를 믿을 것")

for suffix in ("", "_v2", "_v3", "_v4"):
    try:
        target = OUT.with_stem(OUT.stem + suffix)
        doc.save(target)
        print("saved:", target)
        break
    except PermissionError:
        print("잠김(열려 있음):", target.name)
